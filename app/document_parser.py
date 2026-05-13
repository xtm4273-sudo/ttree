"""
文档解析模块 - 支持 PDF/Word/Excel，使用AI智能分段
不依赖固定格式、固定页码、固定编码体系
"""
import os
import sys
import json
import tempfile
import re
from typing import Callable, Optional
from concurrent.futures import ThreadPoolExecutor, as_completed
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
import config
from app.llm_client import get_llm_client


# progress_callback(event, current, total, message)
# event: extracting_text | rule_extract | segmenting_chunks | post_processing
ProgressFn = Optional[Callable[[str, int, int, Optional[str]], None]]


def parse_document(
    file_path: str,
    log_callback=None,
    progress_callback: ProgressFn = None,
) -> list[dict]:
    """
    解析任意格式的询价单文档。

    支持: PDF, Word(.docx), Excel(.xlsx/.xls)

    返回:
    [
        {
            "title": "Crane Service | Shore Crane Hourly Usage",
            "description": "Please provide quotation for crane service...",
            "sfi_code": "1.EH.1.1" 或 None,
            "quantity": 3 或 None,
            "unit": "HR" 或 "",
        },
        ...
    ]
    """
    ext = os.path.splitext(file_path)[1].lower()

    _log = log_callback or (lambda msg: None)
    _prog = progress_callback or (lambda *args: None)

    _log(f"检测到文件格式: {ext}")
    _prog("extracting_text", 0, 1, "正在读取并提取文档文本…")
    if ext == ".pdf":
        raw_text = _extract_pdf(file_path)
    elif ext in (".docx", ".doc"):
        raw_text = _extract_word(file_path)
    elif ext in (".xlsx", ".xls"):
        raw_text = _extract_excel(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

    _log(f"文本提取完成，共 {len(raw_text)} 字符")
    _prog("extracting_text", 1, 1, f"文本已提取（{len(raw_text)} 字符），准备解析条目…")

    # 规则优先：先做结构化抽取，保证条目召回；再用AI补齐未归属内容
    if config.ENABLE_RULE_FIRST and ext in (".xlsx", ".xls"):
        _prog("rule_extract", 0, 1, "正在从表格识别维修条目…")
        rule_items = _rule_segment_excel(file_path)
        if rule_items:
            _log(f"规则抽取命中 {len(rule_items)} 条（Excel结构化优先）")
            _prog("rule_extract", 1, 1, f"已从表格识别 {len(rule_items)} 条")
            _prog("post_processing", 0, 1, "正在整理解析结果…")
            out = _enrich_hierarchy(_deduplicate(rule_items))
            _prog("post_processing", 1, 1, f"解析完成，共 {len(out)} 条")
            return out
        _log("规则抽取未命中，回退AI分段")

    if config.ENABLE_RULE_FIRST:
        _prog("rule_extract", 0, 1, "正在按规则识别条目并审计覆盖率…")
        rule_items, unresolved_lines = _rule_segment_text(raw_text)
        _log(f"规则抽取命中 {len(rule_items)} 条，未归属行 {len(unresolved_lines)} 条")
        _prog("rule_extract", 1, 1, f"规则识别 {len(rule_items)} 条，未归属 {len(unresolved_lines)} 行")

        # 规则结果先入，尽量确保每条都可解释；仅对未归属内容做AI补全，避免全量AI漏项。
        ai_items = []
        if unresolved_lines:
            unresolved_text = "\n".join(unresolved_lines)
            ai_items = _ai_segment(unresolved_text, log_callback=_log, progress_callback=progress_callback)
            _log(f"未归属内容AI补齐 {len(ai_items)} 条")

        merged = _enrich_hierarchy(_deduplicate(rule_items + ai_items))
        _prog("post_processing", 1, 1, f"解析完成，共 {len(merged)} 条")
        return merged

    # 关闭规则优先时，退回全量AI分段
    items = _ai_segment(raw_text, log_callback=_log, progress_callback=progress_callback)
    return items


def _extract_pdf(path: str) -> str:
    """从PDF提取文本（正文+表格），尽可能保留结构行"""
    import pdfplumber

    texts = []
    with pdfplumber.open(path) as pdf:
        for page_idx, page in enumerate(pdf.pages, start=1):
            page_chunks = []

            text = page.extract_text() or ""
            if text.strip():
                page_chunks.append(text)

            # 额外抽取表格，减少仅靠正文OCR导致的漏项
            try:
                tables = page.extract_tables() or []
            except Exception:
                tables = []
            for tbl in tables:
                for row in tbl or []:
                    row_vals = []
                    for cell in row or []:
                        cell_text = str(cell).strip() if cell is not None else ""
                        if cell_text:
                            row_vals.append(cell_text)
                    if row_vals:
                        page_chunks.append(" | ".join(row_vals))

            if page_chunks:
                texts.append(f"=== Page {page_idx} ===\n" + "\n".join(page_chunks))
    return "\n\n".join(texts)


def _extract_word(path: str) -> str:
    """从Word文档提取全部文本"""
    from docx import Document

    doc = Document(path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)

    # 也提取表格内容
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                texts.append(row_text)

    return "\n".join(texts)


def _extract_excel(path: str) -> str:
    """从Excel提取全部文本"""
    import pandas as pd

    dfs = pd.read_excel(path, sheet_name=None, dtype=str)
    texts = []
    for sheet_name, df in dfs.items():
        texts.append(f"=== Sheet: {sheet_name} ===")
        # 转为文本表格
        for _, row in df.iterrows():
            row_text = " | ".join(str(v) for v in row.values if pd.notna(v) and str(v).strip())
            if row_text:
                texts.append(row_text)

    return "\n".join(texts)


def _ai_segment(
    raw_text: str,
    log_callback=None,
    progress_callback: ProgressFn = None,
) -> list[dict]:
    """
    用LLM对原始文本做智能分段，识别出每个独立的维修项目条目。
    不依赖固定格式。
    """
    _log = log_callback or (lambda msg: None)
    _prog = progress_callback or (lambda *args: None)

    # 文本太长时分块处理（较小块降低单次输出截断丢项）
    chunks = _split_text_chunks(raw_text, max_chars=config.PARSE_CHUNK_MAX_CHARS)
    all_items = []

    total_chunks = len(chunks)
    if total_chunks == 0:
        _prog("post_processing", 1, 1, "未从文档中提取到可解析文本")
        return []

    _log(f"文本已分为 {len(chunks)} 个块，开始AI并发解析...")
    _prog(
        "segmenting_chunks",
        0,
        total_chunks,
        f"文档分为 {total_chunks} 段，开始 AI 解析（0/{total_chunks}）…",
    )

    client = get_llm_client()

    # 并发调用LLM解析每个块
    chunk_results = [None] * total_chunks

    def _parse_one(i, chunk):
        items, meta = _segment_chunk_with_optional_resplit(client, chunk, i + 1, total_chunks, _log)
        return i, items, meta

    with ThreadPoolExecutor(max_workers=min(8, total_chunks)) as pool:
        futures = {pool.submit(_parse_one, i, c): i for i, c in enumerate(chunks)}
        done_count = 0
        for future in as_completed(futures):
            i, items, meta = future.result()
            chunk_results[i] = items
            done_count += 1
            extra = f" [补跑]" if meta.get("resplit") else ""
            _log(f"  块 {i+1} 完成: {len(items)} 条 ({done_count}/{total_chunks}){extra}")
            _prog(
                "segmenting_chunks",
                done_count,
                total_chunks,
                f"AI 解析中 {done_count}/{total_chunks} 段…",
            )

    all_items = []
    for items in chunk_results:
        if items:
            all_items.extend(items)

    _prog("post_processing", 0, 1, "正在整理层级与去重…")
    # 层级补全与质量标记（父子关系不完整时强制打标）
    all_items = _enrich_hierarchy(all_items)

    # 去重（同一个条目可能在分块边界重复出现）
    before_dedup = len(all_items)
    all_items = _deduplicate(all_items)
    _log(f"去重: {before_dedup} → {len(all_items)} 条")
    _prog("post_processing", 1, 1, f"解析完成，共 {len(all_items)} 条")

    return all_items


def _annotate_chunk_items(items: list[dict], chunk_idx: int, chunk_text: str, pass_no: int) -> None:
    """为每条解析结果写入溯源字段（就地修改）。"""
    snippet = (chunk_text or "")[:240].replace("\n", " ").strip()
    for it in items:
        it["source_chunk"] = chunk_idx
        it["parse_pass"] = pass_no
        it["raw_text_snippet"] = snippet


def _segment_chunk_with_optional_resplit(
    client,
    text_chunk: str,
    chunk_num: int,
    total_chunks: int,
    log_fn,
) -> tuple[list[dict], dict]:
    """
    先整段解析；若截断/疑似漏项则拆成更小子块重跑并合并（宁可多跑不漏）。
    """
    items, meta = _call_llm_segment_once(client, text_chunk, chunk_num, total_chunks)
    meta = dict(meta)
    meta.setdefault("resplit", False)
    _annotate_chunk_items(items, chunk_num, text_chunk, pass_no=1)

    if not _should_resplit_chunk(text_chunk, items, meta):
        return items, meta

    subs = _split_text_chunks(text_chunk, max_chars=config.PARSE_CHUNK_FALLBACK_CHARS)
    if len(subs) <= 1:
        return items, meta

    merged: list[dict] = []
    for sub in subs:
        sub_items, _ = _call_llm_segment_once(client, sub, chunk_num, total_chunks)
        _annotate_chunk_items(sub_items, chunk_num, sub, pass_no=2)
        merged.extend(sub_items)

    log_fn(
        f"  块 {chunk_num} 补跑拆段: {len(subs)} 子段, 首轮 {len(items)} 条 -> 合并 {len(merged)} 条"
    )
    meta["resplit"] = True
    meta["first_pass_count"] = len(items)
    meta["merged_count"] = len(merged)
    return merged, meta


def _should_resplit_chunk(chunk: str, items: list, meta: dict) -> bool:
    if meta.get("finish_reason") == "length":
        return True
    if meta.get("json_truncated"):
        return True
    if len(items) == 0 and len(chunk.strip()) > 800:
        return True
    # 大块但条目过少：弱启发补跑（召回优先）
    if len(chunk) > int(config.PARSE_CHUNK_MAX_CHARS * 0.85) and len(items) < 3:
        return True
    return False


def _call_llm_segment_once(
    client, text_chunk: str, chunk_num: int, total_chunks: int
) -> tuple[list[dict], dict]:
    """单次 LLM 分段调用，返回 (条目列表, 元信息)。"""
    meta: dict = {"finish_reason": None, "json_truncated": False, "json_repair": False, "error": None}

    prompt = f"""你是船舶维修行业专家。请从以下文本中识别出所有独立的维修/服务项目条目。

## 要求
1. 每个维修项目提取为一条记录
2. 如果有SFI编码就提取，没有就填null
3. 提取标题、描述、数量、单位
4. 忽略目录、页眉页脚、封面等非项目内容
5. **title 仅写客户询价项目名称本身**：不得包含数量、单位；不得用半角()或全角（）在标题里追加补充说明（补充说明一律写入 description）
6. 描述要保留所有关键技术细节和多行任务内容，用 | 分隔不同任务行
7. 识别SFI层级关系，如 3.ME.11 是父级，3.ME.11.1 是子级，在子级上标注 parent_sfi
8. 识别区间写法（如 4.AE.1.1-4.AE.1.6），提取为一个条目并标注 is_range: true
9. 严禁只抽“父级总标题”而漏掉其子项：若文本里存在子项，必须逐条输出子项
10. 若你无法确认子项是否完整，给该条添加 quality_flags: ["possible_missing_children"]
11. **父级与明细都要输出**：凡出现父级标题（如章节/SFI父级总标题），必须单独输出一条父级记录；其下每个子项再各输出一条，不得把父级与子项合并成一条。

## 文本内容（第{chunk_num}/{total_chunks}块）
{text_chunk}

## 输出格式
返回JSON数组，每项格式如下：
```json
[
  {{
    "title": "仅项目标题本身，不含数量、单位与括注",
    "description": "详细描述（保留所有技术要求，多行用 | 分隔）",
    "sfi_code": "1.EH.1.1" 或 null,
    "parent_sfi": "父级SFI编码" 或 null,
    "is_range": true 或 false,
    "quality_flags": ["possible_missing_children"] 或 [],
    "quantity": 数字 或 null,
    "unit": "PC/SET/DAY/HR/LOT等" 或 ""
  }}
]
```
只返回JSON数组，不要其他文字。如果这段文本中没有维修项目，返回空数组 []。"""

    try:
        response = client.chat.completions.create(
            model=config.LLM_MODEL,
            messages=[{"role": "user", "content": prompt}],
            temperature=0.1,
            max_tokens=config.PARSE_MAX_TOKENS,
        )
        choice = response.choices[0]
        meta["finish_reason"] = getattr(choice, "finish_reason", None)
        content = (choice.message.content or "").strip()
        json_str = _extract_json_array(content)
        if json_str and not json_str.rstrip().endswith("]"):
            meta["json_truncated"] = True
        if not json_str:
            return [], meta

        items_raw, repaired = _safe_parse_json_array_ex(json_str)
        meta["json_repair"] = repaired

        valid_items = []
        for item in items_raw:
            if isinstance(item, dict) and item.get("title"):
                qty = item.get("quantity")
                unt = str(item.get("unit", ""))
                t = _clean_enquiry_title_for_customer(str(item.get("title", "")), qty, unt)
                valid_items.append({
                    "title": t,
                    "description": str(item.get("description", "")),
                    "sfi_code": item.get("sfi_code"),
                    "parent_sfi": item.get("parent_sfi"),
                    "is_range": item.get("is_range", False),
                    "quality_flags": list(item.get("quality_flags", [])) if isinstance(item.get("quality_flags"), list) else [],
                    "quantity": qty,
                    "unit": unt,
                })
        return valid_items, meta

    except Exception as e:
        meta["error"] = str(e)
        print(f"[文档解析] 第{chunk_num}块AI分段失败: {e}")

    return [], meta


def _split_text_chunks(text: str, max_chars: int | None = None) -> list[str]:
    """
    将长文本切分为合适大小的块。
    尽量在段落边界切分，避免切断一个条目。
    """
    if max_chars is None:
        max_chars = config.PARSE_CHUNK_MAX_CHARS
    if len(text) <= max_chars:
        return [text]

    chunks = []
    paragraphs = text.split("\n\n")
    current_chunk = ""

    for para in paragraphs:
        if len(current_chunk) + len(para) + 2 > max_chars:
            if current_chunk:
                chunks.append(current_chunk)
            current_chunk = para
        else:
            current_chunk += "\n\n" + para if current_chunk else para

    if current_chunk:
        chunks.append(current_chunk)

    # 如果单个段落就超长，强制切分
    final_chunks = []
    for chunk in chunks:
        if len(chunk) > max_chars:
            for i in range(0, len(chunk), max_chars):
                final_chunks.append(chunk[i:i + max_chars])
        else:
            final_chunks.append(chunk)

    return final_chunks


def _extract_json_array(text: str) -> str | None:
    """从LLM返回文本中提取JSON数组"""
    import re
    # 尝试 ```json ... ```
    match = re.search(r'```json\s*([\s\S]*?)\s*```', text)
    if match:
        return match.group(1)
    # 尝试找 [ ... ]
    match = re.search(r'\[[\s\S]*\]', text)
    if match:
        return match.group(0)
    # 可能被截断，只有 [ 开头没有 ] 结尾
    match = re.search(r'\[[\s\S]*', text)
    if match:
        return match.group(0)
    return None


def _safe_parse_json_array_ex(json_str: str) -> tuple[list, bool]:
    """解析JSON数组，支持截断修复。返回 (列表, 是否使用了截断修复)。"""
    try:
        result = json.loads(json_str)
        if isinstance(result, list):
            return result, False
    except json.JSONDecodeError:
        pass

    for i in range(len(json_str) - 1, 0, -1):
        if json_str[i] == '}':
            attempt = json_str[:i+1] + ']'
            try:
                result = json.loads(attempt)
                if isinstance(result, list):
                    return result, True
            except json.JSONDecodeError:
                continue

    return [], False


def _safe_parse_json_array(json_str: str) -> list:
    """兼容旧接口：仅返回列表。"""
    items, _ = _safe_parse_json_array_ex(json_str)
    return items


def _clean_enquiry_title_for_customer(title: str, quantity, unit: str) -> str:
    """
    客户可见「询价标题」：去掉括注、去掉已单独抽取的数量与单位（仅尾部清洗）。
    数量/单位仍保留在结构化字段中。
    """
    raw = re.sub(r"\s+", " ", (title or "").strip())
    if not raw:
        return ""

    s = raw
    tail_half = re.compile(r"\s*\([^)]{0,120}\)\s*$")
    tail_full = re.compile(r"\s*（[^）]{0,120}）\s*$")
    for _ in range(24):
        ns = tail_half.sub("", s)
        if ns != s:
            s = ns
            continue
        ns = tail_full.sub("", s)
        if ns != s:
            s = ns
            continue
        break

    u = (unit or "").strip().upper()
    q = quantity
    if q is not None and str(q).strip() != "":
        try:
            qf = float(q)
            if qf.is_integer():
                q_strs = {str(int(qf)), str(qf)}
            else:
                q_strs = {str(qf), str(q).strip()}
        except (TypeError, ValueError):
            q_strs = {str(q).strip()}
        for qs in list(q_strs):
            if "." in qs and qs.replace(".", "", 1).isdigit():
                try:
                    qf2 = float(qs)
                    if qf2.is_integer():
                        q_strs.add(str(int(qf2)))
                except ValueError:
                    pass

        variants = sorted({x for x in q_strs if x}, key=len, reverse=True)
        if u:
            u_esc = re.escape(u)
            for qs in variants:
                qe = re.escape(qs)
                for pat in (
                    rf"(?i)[\s,|/\-]*(?:x\s*)?{qe}[\s,|/\-]*{u_esc}\s*$",
                    rf"(?i)[\s,|/\-]*{qe}[\s]*{u_esc}\s*$",
                ):
                    s = re.sub(pat, "", s)
        for qs in variants:
            qe = re.escape(qs)
            s = re.sub(rf"(?i)[\s,|/\-]+(?:x\s*)?{qe}\s*$", "", s)

    s = re.sub(r"\s+", " ", s).strip(" -|/,，")
    if not s:
        return raw[:200]
    return s[:200]


def _deduplicate(items: list[dict]) -> list[dict]:
    """最小去重：仅删除 SFI+标题+描述+数量+单位 完全相同的重复项。"""
    seen_keys = set()
    unique_items = []

    for item in items:
        key = _dedup_key_full(item)
        if key not in seen_keys:
            seen_keys.add(key)
            unique_items.append(item)

    return unique_items


def _dedup_key_full(item: dict) -> tuple:
    sfi = _normalize_sfi(item.get("sfi_code")) or ""
    title = re.sub(r"\s+", " ", str(item.get("title", "")).strip())
    desc = re.sub(r"\s+", " ", str(item.get("description", "")).strip())
    qty = "" if item.get("quantity") is None else str(item.get("quantity"))
    unit = re.sub(r"\s+", "", str(item.get("unit", "")).strip().upper())
    return (sfi, title, desc, qty, unit)


def _rule_segment_text(raw_text: str) -> tuple[list[dict], list[str]]:
    """
    通用规则抽取（PDF/Word/非结构化Excel文本）：
    - 优先按SFI行起项；
    - 对明显序号条目做兜底；
    - 返回未归属行，供AI二次补齐。
    """
    unit_candidates = {
        "PC", "PCS", "SET", "DAY", "DAYS", "HR", "HRS", "LOT", "M", "M2", "M3", "KG", "EA", "LIFT"
    }
    sfi_regex = re.compile(r"(?<!\w)(\d{1,2}\s*\.\s*[A-Z]{1,4}\s*\.\s*\d+(?:\s*\.\s*\d+)*)")
    numbered_regex = re.compile(r"^\s*\d+[\.\)]\s+(.+)$")
    qty_regex = re.compile(r"^\d+(?:\.\d+)?$")

    lines = []
    for raw in raw_text.splitlines():
        line = re.sub(r"\s+", " ", raw).strip()
        if line:
            lines.append(line)

    consumed = [False] * len(lines)
    items = []
    current = None

    def _finalize_current():
        nonlocal current
        if not current:
            return
        title = (current.get("title") or "").strip()
        if len(title) < 3:
            current = None
            return
        current["description"] = " | ".join(current.get("desc_parts", []))[:1200]
        current.pop("desc_parts", None)
        current["title"] = _clean_enquiry_title_for_customer(
            title,
            current.get("quantity"),
            str(current.get("unit") or ""),
        )
        items.append(current)
        current = None

    def _parse_qty_unit(tokens: list[str]) -> tuple[float | int | None, str]:
        quantity = None
        unit = ""
        for tok in tokens:
            t = tok.strip().upper()
            if not unit and t in unit_candidates:
                unit = t
                continue
            if quantity is None and qty_regex.match(tok.strip()):
                try:
                    q = float(tok.strip())
                    quantity = int(q) if q.is_integer() else q
                except ValueError:
                    pass
        return quantity, unit

    for i, line in enumerate(lines):
        sfi_match = sfi_regex.search(line)
        if sfi_match:
            _finalize_current()
            sfi = _normalize_sfi(sfi_match.group(1))
            after = line[sfi_match.end():].strip(" -:|")
            tokens = re.split(r"[|/\-,\s]+", line)
            quantity, unit = _parse_qty_unit(tokens)
            cleaned = _clean_enquiry_title_for_customer(
                (after[:200] if after else ""), quantity, str(unit or "")
            )
            title = cleaned or (f"SFI {sfi}" if sfi else "")
            current = {
                "title": title,
                "description": "",
                "sfi_code": sfi,
                "parent_sfi": _get_parent_sfi(sfi),
                "is_range": bool(re.search(r"\bTO\b|\-.*\.", line.upper())),
                "quality_flags": [],
                "quantity": quantity,
                "unit": unit,
                "desc_parts": [],
            }
            consumed[i] = True
            continue

        # 无SFI但像编号条目，做兜底
        num_match = numbered_regex.match(line)
        if num_match and len(num_match.group(1)) > 8:
            _finalize_current()
            raw_title = num_match.group(1)[:200]
            tokens = re.split(r"[|/\-,\s]+", line)
            quantity, unit = _parse_qty_unit(tokens)
            title = _clean_enquiry_title_for_customer(
                raw_title, quantity, str(unit or "")
            ) or raw_title.strip()
            current = {
                "title": title,
                "description": "",
                "sfi_code": None,
                "parent_sfi": None,
                "is_range": False,
                "quality_flags": ["missing_sfi_hierarchy_uncertain"],
                "quantity": quantity,
                "unit": unit,
                "desc_parts": [],
            }
            consumed[i] = True
            continue

        # 作为当前条目的延续描述
        if current and len(line) >= 3:
            upper_line = line.upper()
            # 过滤常见页眉页脚噪音
            if not re.match(r"^(PAGE|P\.?\s*\d+|NO\.?|DATE|VESSEL|INQUIRY)\b", upper_line):
                if not current.get("title") or current["title"].startswith("SFI "):
                    current["title"] = line[:200]
                else:
                    current["desc_parts"].append(line[:300])
                consumed[i] = True

    _finalize_current()

    unresolved_lines = []
    for i, line in enumerate(lines):
        if consumed[i]:
            continue
        # 保留有信息量的未归属行，供二次AI补齐
        if len(line) >= 10 and re.search(r"[A-Za-z]", line):
            unresolved_lines.append(line)

    # 避免超长文本导致AI再次截断，保留前后关键区域
    if len(unresolved_lines) > 300:
        unresolved_lines = unresolved_lines[:180] + unresolved_lines[-120:]

    for it in items:
        it.setdefault("source_chunk", 0)
        it.setdefault("parse_pass", 0)
        it.setdefault("raw_text_snippet", "")

    return items, unresolved_lines


def _rule_segment_excel(path: str) -> list[dict]:
    """
    Excel规则抽取：
    - 优先识别含SFI编码的行；
    - 按行提取标题/描述/数量/单位；
    - 仅用于“能直接命中结构化信息”的场景，避免不必要LLM调用。
    """
    import pandas as pd

    sfi_pattern = re.compile(r"^\d{1,2}\.[A-Z]{1,4}\.\d+(?:\.\d+)*$")
    unit_candidates = {"PC", "PCS", "SET", "DAY", "DAYS", "HR", "HRS", "LOT", "M", "M2", "M3", "KG", "EA"}

    try:
        sheets = pd.read_excel(path, sheet_name=None, dtype=str)
    except Exception:
        return []

    items = []
    for _, df in sheets.items():
        for _, row in df.iterrows():
            cells = []
            for v in row.values:
                if pd.notna(v):
                    text = str(v).strip()
                    if text:
                        cells.append(text)
            if not cells:
                continue

            sfi = None
            sfi_idx = -1
            for i, c in enumerate(cells):
                candidate = c.upper().replace(" ", "")
                if sfi_pattern.match(candidate):
                    sfi = candidate
                    sfi_idx = i
                    break
            if not sfi:
                continue

            title = ""
            if sfi_idx + 1 < len(cells):
                title = cells[sfi_idx + 1]
            if not title:
                for c in cells:
                    if c.upper().replace(" ", "") != sfi and len(c) >= 4:
                        title = c
                        break
            if not title:
                continue

            quantity = None
            unit = ""
            description_parts = []
            for i, c in enumerate(cells):
                if i == sfi_idx:
                    continue
                c_clean = c.strip()
                c_upper = c_clean.upper()

                if c_upper in unit_candidates and not unit:
                    unit = c_upper
                    continue

                if quantity is None and re.match(r"^\d+(\.\d+)?$", c_clean):
                    try:
                        q = float(c_clean)
                        quantity = int(q) if q.is_integer() else q
                        continue
                    except ValueError:
                        pass

                if c_clean != title:
                    description_parts.append(c_clean)

            title = _clean_enquiry_title_for_customer(title, quantity, unit)
            items.append({
                "title": title,
                "description": " | ".join(description_parts),
                "sfi_code": sfi,
                "parent_sfi": _get_parent_sfi(sfi),
                "is_range": False,
                "quality_flags": [],
                "quantity": quantity,
                "unit": unit,
                "source_chunk": 0,
                "parse_pass": 0,
                "raw_text_snippet": "",
            })

    return items


def _normalize_sfi(sfi: str | None) -> str | None:
    if not sfi:
        return None
    value = str(sfi).strip().upper()
    value = re.sub(r"\s+", "", value)
    if not value:
        return None
    # 仅保留常见SFI字符
    value = re.sub(r"[^A-Z0-9\.\-]", "", value)
    return value or None


def _get_parent_sfi(sfi: str | None) -> str | None:
    if not sfi or "." not in sfi:
        return None
    parts = sfi.split(".")
    if len(parts) <= 1:
        return None
    return ".".join(parts[:-1])


def _sfi_depth(sfi: str | None) -> int:
    if not sfi:
        return 0
    return len([p for p in sfi.split(".") if p])


def _enrich_hierarchy(items: list[dict]) -> list[dict]:
    """
    层级增强：
    1) 规范化SFI/parent_sfi并自动补parent_sfi；
    2) 对“仅父级、未识别子级”打标，供后续匹配阶段强制降级。
    """
    if not items:
        return items

    # 先标准化
    normalized = []
    for item in items:
        sfi = _normalize_sfi(item.get("sfi_code"))
        parent = _normalize_sfi(item.get("parent_sfi"))
        inferred_parent = _get_parent_sfi(sfi)
        if not parent and inferred_parent:
            parent = inferred_parent

        flags = item.get("quality_flags") or []
        if not isinstance(flags, list):
            flags = []

        normalized.append({
            **item,
            "sfi_code": sfi,
            "parent_sfi": parent,
            "quality_flags": flags,
        })

    # 统计已识别子项
    all_sfi = {it.get("sfi_code") for it in normalized if it.get("sfi_code")}
    has_children = set()
    for sfi in all_sfi:
        p = _get_parent_sfi(sfi)
        while p:
            has_children.add(p)
            p = _get_parent_sfi(p)

    # 给风险条目标记
    for it in normalized:
        sfi = it.get("sfi_code")
        flags = set(it.get("quality_flags") or [])
        depth = _sfi_depth(sfi)
        desc = (it.get("description") or "").strip()

        # 规则：是父级(<=3段)、无子项、且描述很短/为空 => 高风险
        if sfi and depth <= 3 and sfi not in has_children and not it.get("is_range"):
            if len(desc) < 20:
                flags.add("parent_without_children")

        # 没有SFI但被模型提示可能缺子项
        if not sfi and "possible_missing_children" in flags:
            flags.add("missing_sfi_hierarchy_uncertain")

        it["quality_flags"] = sorted(flags)

    return normalized
