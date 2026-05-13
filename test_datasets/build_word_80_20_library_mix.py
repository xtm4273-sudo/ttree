"""
生成 Word 询价单：约 80% 条目 SFI+描述取自当前 data/craft_library.json（库内），
约 20% 为虚构 SFI + 非坞修语义（库内无对应标准项），用于测向量置信度与待确认行为。

说明：解析器 SFI 行须以数字开头（见 document_parser._rule_segment_text），故「库外」
条目使用 99.XX.* / 98.IT.* 等形式，而非 Z.NE.*。

输出：与本脚本同目录下的 sample_enquiry_80_20_library_mix.docx
"""
from __future__ import annotations

import json
import os
import re

from docx import Document
from docx.shared import Pt

IN_LIBRARY_COUNT = 20
OFF_LIBRARY_COUNT = 5

_REPO_ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
_CRAFT_JSON = os.path.join(_REPO_ROOT, "data", "craft_library.json")


def _out_dir() -> str:
    return os.path.dirname(os.path.abspath(__file__))


def _pick_in_library_items(craft_items: list[dict], n: int) -> list[dict]:
    """选叶子级、无斜杠的 SFI，按编码排序后取前 n 条，保证可复现。"""
    eligible: list[dict] = []
    for it in craft_items:
        sfi = (it.get("sfi_code") or "").strip().upper()
        if not sfi or "/" in sfi:
            continue
        if sfi.count(".") < 3:
            continue
        if not re.match(r"^\d{1,2}\.[A-Z]{1,4}\.\d+(?:\.\d+)*$", sfi):
            continue
        title = (it.get("title") or "").strip()
        if len(title) < 4:
            continue
        eligible.append(it)

    eligible.sort(key=lambda x: (x.get("sfi_code") or ""))
    if len(eligible) < n:
        raise RuntimeError(f"craft_library 可抽取叶子条目不足 {n}，当前 {len(eligible)}")
    return eligible[:n]


def _zh_note_for(title: str, detail: str) -> str:
    """简短中文备注，便于人读；向量仍以英文 title/detail 为主。"""
    t = (title or "")[:80]
    d = (detail or "")[:60].replace("|", " ")
    if d:
        return f"{t}（{d}…）"
    return t


def build_docx(path: str) -> None:
    with open(_CRAFT_JSON, "r", encoding="utf-8") as f:
        craft_items = json.load(f)

    picked = _pick_in_library_items(craft_items, IN_LIBRARY_COUNT)

    off_rows: list[tuple[str, str, str, str]] = [
        (
            "99.XX.1.1",
            "低轨卫星通信桅杆天线罩 GRP 修补、雷击点补强及岸端 EMC 抽检配合",
            "1",
            "LOT",
        ),
        (
            "98.IT.9.9",
            "定制 Python：历史 AIS 报文清洗并迁移至岸端 SQL Server（纯 IT 服务）",
            "1",
            "JOB",
        ),
        (
            "97.NF.1.1",
            "全船 NFT 数字艺术走廊策展与灯光编程（与船舶维修无关）",
            "1",
            "DAY",
        ),
        (
            "96.ML.2.2",
            "船上 Wi-Fi 行为分析模型训练与边缘推理容器部署（机器学习外包）",
            "1",
            "SET",
        ),
        (
            "95.QA.3.3",
            "第三方渗透测试与漏洞扫描（网络安全咨询，非船级检验）",
            "1",
            "LOT",
        ),
    ]
    if len(off_rows) != OFF_LIBRARY_COUNT:
        raise ValueError("off_rows 数量须与 OFF_LIBRARY_COUNT 一致")

    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("万邦船舶 — 询价单 MOCK（约 80% 库内 / 20% 库外）")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(
        f"船名 MV CONFIDENCE MIX  |  港口 SHANGHAI  |  "
        f"共 {IN_LIBRARY_COUNT + OFF_LIBRARY_COUNT} 条："
        f"{IN_LIBRARY_COUNT} 条取自 craft_library.json；"
        f"{OFF_LIBRARY_COUNT} 条为虚构 SFI（库外语义）。"
    )

    doc.add_paragraph("—— 以下为工艺库内条目（规则分段 SFI 行）——")

    for it in picked:
        sfi = (it.get("sfi_code") or "").strip()
        title = (it.get("title") or "").strip()
        detail = (it.get("detail") or "").strip()
        unit = (it.get("unit") or "LOT").strip().upper() or "LOT"
        note = _zh_note_for(title, detail)
        qty = it.get("qty_template")
        try:
            qstr = str(int(qty)) if qty is not None and str(qty).strip().isdigit() else "1"
        except (TypeError, ValueError):
            qstr = "1"
        doc.add_paragraph(f"{sfi} {note}，数量 {qstr} {unit}")

    doc.add_paragraph("—— 以下为工艺库外条目（虚构 SFI，库内无对应项）——")

    for sfi, desc, qty, unit in off_rows:
        doc.add_paragraph(f"{sfi} {desc}，数量 {qty} {unit}")

    doc.save(path)


def main() -> None:
    out = os.path.join(_out_dir(), "sample_enquiry_80_20_library_mix.docx")
    build_docx(out)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    main()
