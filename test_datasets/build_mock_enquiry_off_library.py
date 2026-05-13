"""
生成「库内 + 明显库外/虚构」混排询价样例，用于观察匹配表现。

说明（与 match_engine._build_result 一致）：
- `is_new_item=True` 仅在向量检索得到 **零条候选** 时成立；FAISS 在索引非空时几乎总会返回
  Top-K 邻居，因此多数「工艺库里没有对应标准项」的询价仍会走候选 + 置信度路径。
- 更常见现象是：`needs_human_review=True`、`review_status=PENDING_REVIEW`、置信度偏低、
  Top1 工艺条目与客户描述明显不符（需人工确认或学习入库）。
- 若开启 config.ENABLE_AUTO_SUGGEST_NEW_ENTRY，仅在真正零候选时会由 LLM 生成 suggested_entry；
  默认配置下该开关为 False。

输出：与本脚本同目录下的 sample_enquiry_off_library.xlsx / .docx
"""
from __future__ import annotations

import os

from openpyxl import Workbook
from docx import Document
from docx.shared import Pt


def _out_dir() -> str:
    return os.path.dirname(os.path.abspath(__file__))


def build_excel(path: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "RepairList"
    ws.append(["SFI", "工作描述 / Description", "数量", "单位"])
    rows = [
        # --- 工艺库内常见 SFI（craft_library.json 中存在），便于对照「正常命中」---
        ("1.MS.1.1", "船底塞拆装换新垫料/水泥（与模板工艺一致）", "16", "PC"),
        ("1.EH.1.1", "岸吊/浮吊按小时使用（与库内 CRANE SERVICE 语义接近）", "8", "HR"),
        # --- 虚构 SFI + 工艺库中不存在的业务描述：预期弱匹配 / 待人工确认 ---
        (
            "Z.NE.1.1",
            "低轨卫星通信桅杆天线罩 GRP 裂纹修补、雷击点补强及岸端 EMC 抽检配合",
            "1",
            "LOT",
        ),
        (
            "X.IT.9.9",
            "定制 Python 脚本：历史 AIS 报文批量清洗并迁移至岸端 SQL Server（纯 IT 服务）",
            "1",
            "JOB",
        ),
        (
            "9.XX.1.1",
            "全船 NFT 艺术走廊策展与灯光编程（与船舶维修无关的占位描述）",
            "1",
            "DAY",
        ),
    ]
    for r in rows:
        ws.append(list(r))
    wb.save(path)


def build_word(path: str) -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("万邦船舶 — 询价单 MOCK（含工艺库外语义条目）")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("船名 MV MOCK OFF-LIB  |  港口 NINGBO  |  用途：匹配回归测试")
    doc.add_paragraph(
        "含两条与工艺库接近的正常维修项，以及三条虚构 SFI / 非坞修业务，用于观察向量+门控+置信度表现。"
    )

    doc.add_paragraph("1.MS.1.1 船底塞拆装，按实际数量 16 PC")
    doc.add_paragraph("1.EH.1.1 岸吊服务 8 HR")

    doc.add_paragraph(
        "Z.NE.1.1 低轨卫星通信桅杆天线罩 GRP 修补及 EMC 抽检配合，数量 1 LOT"
    )
    doc.add_paragraph(
        "X.IT.9.9 AIS 报文迁移脚本开发，岸端数据库对接，1 JOB"
    )

    doc.add_paragraph("表格行：")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "SFI"
    hdr[1].text = "Description"
    hdr[2].text = "Qty"
    hdr[3].text = "Unit"
    for sfi, desc, qty, unit in [
        ("9.XX.1.1", "NFT 艺术走廊策展与灯光编程", "1", "DAY"),
    ]:
        row = table.add_row().cells
        row[0].text = sfi
        row[1].text = desc
        row[2].text = qty
        row[3].text = unit

    doc.save(path)


def main() -> None:
    d = _out_dir()
    xlsx = os.path.join(d, "sample_enquiry_off_library.xlsx")
    docx = os.path.join(d, "sample_enquiry_off_library.docx")
    build_excel(xlsx)
    build_word(docx)
    print(f"Wrote: {xlsx}")
    print(f"Wrote: {docx}")


if __name__ == "__main__":
    main()
