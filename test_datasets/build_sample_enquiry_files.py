"""
生成船舶维修询价样例（Excel / Word），用于端到端测试解析与匹配。
Excel：每行含标准 SFI 单元格，可走 _rule_segment_excel 规则优先路径。
Word：段落 + 表格混合，可走 _rule_segment_text（SFI 行起项）。
"""
from __future__ import annotations

import os

from openpyxl import Workbook
from docx import Document
from docx.shared import Pt


def _out_dir() -> str:
    return os.path.dirname(os.path.abspath(__file__))


def build_excel(path: str) -> None:
    wb = Workbook()
    ws = wb.active
    ws.title = "RepairList"
    # 表头（无 SFI，规则行会跳过）
    ws.append(["SFI", "工作描述 / Description", "数量", "单位"])
    rows = [
        ("1.EH.1.1", "锚机解体检查、清洗、复装及负荷试验", "1", "LOT"),
        ("2.CH.1.2", "甲板克令吊钢丝绳更换，含旧绳回收", "2", "PCS"),
        ("3.HV.2.1", "压载泵解体检查、叶轮探伤、机械密封更换", "1", "SET"),
        ("4.PI.3.2", "主空压机进气滤器、油分滤芯整套更换", "4", "PCS"),
        ("5.EL.1.1", "应急发电机年度检验及负载试验（按船级社要求）", "1", "DAY"),
    ]
    for r in rows:
        ws.append(list(r))
    wb.save(path)


def build_word(path: str) -> None:
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("万邦船舶 — 维修询价单（样例 / SAMPLE）")
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph("船名 MV SAMPLE STAR  |  港口 SHANGHAI  |  计划进坞 2026-Q3")
    doc.add_paragraph(
        "下列项目请按现行工艺与物料标准报价；含人工、备件、试验及报告。"
    )

    # SFI 起项的独立行（规则分段可识别）
    doc.add_paragraph(
        "1.EH.1.1 锚机整体检修 — 含刹车带测量、液压马达密封更换，数量 1 LOT"
    )
    doc.add_paragraph(
        "2.CH.1.2 克令吊回转轴承润滑与钢丝绳检查，钢丝绳若需更换另列备件。"
    )
    doc.add_paragraph("3.HV.2.1 压载系统阀门液压试验 2 HR")

    doc.add_paragraph("表格汇总（与 Excel 场景对照）：")
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    hdr[0].text = "SFI"
    hdr[1].text = "Description"
    hdr[2].text = "Qty"
    hdr[3].text = "Unit"
    for sfi, desc, qty, unit in [
        ("4.PI.3.2", "主空压机保养包（滤芯套装）", "1", "SET"),
        ("5.EL.1.1", "应急配电板绝缘测试及继电器校验", "1", "DAY"),
    ]:
        row = table.add_row().cells
        row[0].text = sfi
        row[1].text = desc
        row[2].text = qty
        row[3].text = unit

    doc.save(path)


def main() -> None:
    d = _out_dir()
    xlsx = os.path.join(d, "sample_enquiry_excel.xlsx")
    docx = os.path.join(d, "sample_enquiry_word.docx")
    build_excel(xlsx)
    build_word(docx)
    print(f"Wrote: {xlsx}")
    print(f"Wrote: {docx}")


if __name__ == "__main__":
    main()
